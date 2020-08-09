import datetime
import os
import sqlite3
import time

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from openpyxl.styles import Border, Side
from PIL import Image

from libs.exception import AppException
from libs.helper import md5


def read_excel(filename, start_row, title):
    ws = load_workbook(filename).active

    ws_title = ws.cell(1, 1).value
    if ws_title != title:
        raise AppException('导入文件的标题错误，请使用正确的模板导入')

    col = ws.max_column
    row = ws.max_row

    data = []
    for row_idx in range(start_row, row + 1):
        try:
            tmp = []
            for col_idx in range(2, col + 1):
                cell = ws.cell(row_idx, col_idx)
                tmp.append(str(cell.value) if cell.value else None)
        except IndexError:
            raise AppException('导入数据错误，请使用模板导入数据')
        data.append(tmp)
    return data


def save_excel(template_filename, start_row, data, filename):
    wb = load_workbook(template_filename)
    ws = wb.active

    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))

    for row_idx in range(start_row, len(data) + start_row):
        for col_idx in range(1, len(data[0]) + 1):
            cell = ws.cell(row_idx, col_idx)
            cell.value = data[row_idx - start_row][col_idx - 1]
            cell.border = thin_border

    wb.save(filename)


def upload_file(filename, is_pic=False):
    if not os.path.exists('file/'):
        os.makedirs('file/')
    if is_pic:
        try:
            image = Image.open(filename)
        except Exception:
            raise AppException('图片无法打开')
        filename = 'tmp.png'
        image.save(filename)

    suffix = filename.rsplit('.', 1)
    if len(suffix) == 1:
        raise AppException('文件没有后缀')
    suffix = suffix[1]
    with open(filename, "rb") as f:
        raw = f.read()
    if is_pic:
        os.remove(filename)
    filename = 'file/{}'.format(md5(raw) + "." + suffix)
    with open(filename, "wb") as f:
        f.write(raw)
    return filename


def download_file(from_filename, to_filename):
    try:
        with open(from_filename, "rb") as f:
            raw = f.read()
    except OSError:
        raise AppException('打开原始文件失败，文件可能已经丢失')
    with open(to_filename, "wb") as f:
        f.write(raw)


def save_word(filename, title, data, pic=False, ty_data=None):  # noqa: C901
    if ty_data is None:
        ty_data = []
    document = Document()
    document.styles['Normal'].font.name = u'楷体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(25)
    title_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    pic_row = 5
    if pic:
        row = (len(data) + pic_row) // 2
    else:
        row = (len(data) + 1) // 2
    table = document.add_table(rows=row, cols=4, style='Table Grid')
    table.style.font.size = Pt(16)
    if pic:
        table.cell(0, 2).merge(table.cell(pic_row - 1, 2))
        table.cell(0, 3).merge(table.cell(pic_row - 1, 3))

    row_idx = 0
    col_idx = 0
    for k, v in data.items():
        if k == '照片':
            continue
        cell_k = table.cell(row_idx, col_idx)
        cell_k.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_v = table.cell(row_idx, col_idx + 1)
        cell_v.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_k.text = k
        cell_v.text = v

        if (pic and row_idx < pic_row) or col_idx == 2:
            row_idx += 1
            col_idx = 0
        else:
            col_idx = 2
    if pic:
        cell_k = table.cell(0, 2)
        cell_k.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_v = table.cell(0, 3)
        cell_v.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_k.text = '照片'
        try:
            with open(data['照片'], 'rb'):
                pass
            cell_v.add_paragraph().add_run().add_picture(data['照片'], width=Inches(1.5))
        except OSError:
            pass

    if ty_data:
        document.add_paragraph()
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run('团员')
        title_run.font.size = Pt(20)
        title_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        row = len(ty_data) + 1
        col = len(ty_data[0])

        table = document.add_table(rows=row, cols=col, style='Table Grid')
        idx = 0
        for k in ty_data[0].keys():
            cell = table.cell(0, idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = k
            idx += 1

        for row_idx in range(1, row):
            idx = 0
            for v in ty_data[row_idx - 1].values():
                cell = table.cell(row_idx, idx)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = v
                idx += 1

    document.save(filename)


def backup_database(filename):
    sqlite3.connect('database.db').backup(sqlite3.connect(filename))
    print('备份成功', filename)


def auto_backup():
    if not os.path.exists('backup/'):
        os.makedirs('backup/')
    while 1:
        curr_time = datetime.datetime.now()
        curr_time = curr_time.strftime('%Y-%m-%d-%H-%M-%S')
        filename = 'backup/auto_backup-{}.db'.format(curr_time)
        backup_database(filename)
        time.sleep(3600)
